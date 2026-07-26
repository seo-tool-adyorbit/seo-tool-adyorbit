"""
PDF Service — LibreOffice + python-docx + PyMuPDF Pipeline
============================================================
Flow: HTML → python-docx DOCX → LibreOffice headless PDF → PyMuPDF optimize
Fallback: ReportLab with Arial Unicode & FontAwesome icon rendering
Output: 100% searchable, selectable, text-based PDF (Google Docs quality)
"""

import io
import os
import re
import subprocess
import tempfile
import shutil
from pathlib import Path
from typing import Optional
from bs4 import BeautifulSoup, NavigableString, Tag

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import fitz  # PyMuPDF


# ─────────────────────────────────────────────────────────────
# FontAwesome Class → Unicode Character Mapping
# ─────────────────────────────────────────────────────────────

FA_CLASS_TO_UNICODE = {
    'fa-phone': '☎', 'fa-phone-alt': '☎', 'fa-phone-volume': '☎',
    'fa-headset': '☎', 'fa-mobile': '📱', 'fa-mobile-alt': '📱',
    'fa-check': '✔', 'fa-check-circle': '✔', 'fa-circle-check': '✔',
    'fa-times': '✘', 'fa-xmark': '✘', 'fa-ban': '✘',
    'fa-star': '★', 'fa-star-half': '★',
    'fa-heart': '♥', 'fa-thumbs-up': '👍', 'fa-thumbs-down': '👎',
    'fa-plane': '✈', 'fa-plane-departure': '✈', 'fa-plane-arrival': '✈',
    'fa-bolt': '⚡', 'fa-fire': '🔥', 'fa-rocket': '🚀',
    'fa-envelope': '✉', 'fa-email': '✉',
    'fa-map-marker': '📍', 'fa-location-dot': '📍',
    'fa-clock': '⏰', 'fa-calendar': '📅',
    'fa-user': '👤', 'fa-users': '👥',
    'fa-arrow-right': '➔', 'fa-arrow-left': '⬅', 'fa-arrow-up': '⬆', 'fa-arrow-down': '⬇',
    'fa-circle': '●', 'fa-square': '■',
    'fa-info': 'ℹ', 'fa-info-circle': 'ℹ',
    'fa-warning': '⚠', 'fa-triangle-exclamation': '⚠',
    'fa-globe': '🌐', 'fa-wifi': '📶',
    'fa-dollar': '$', 'fa-dollar-sign': '$',
    'fa-cog': '⚙', 'fa-gear': '⚙', 'fa-gears': '⚙',
    'fa-cube': '🧊', 'fa-folder-open': '📂', 'fa-code': '</>',
    'fa-wand-magic-sparkles': '✨',
}

# Supplementary Plane Emojis -> Clean BMP Vectors
EMOJI_PATTERNS = [
    # Airplanes / Flights -> ✈ (\u2708)
    (re.compile(r'[\U0001F6EB\U0001F6EC\U0001F6E9\U0001F6E8\u2708][\uFE0F\uFE0E]?'), '✈'),
    # Phones / Helpline -> ☎ (\u260E)
    (re.compile(r'[\U0001F4DE\U0001F4F1\U0001F4F2\U0001F919\u260E\u260F][\uFE0F\uFE0E]?'), '☎'),
    # Checks -> ✔ (\u2714)
    (re.compile(r'[\u2705\u2714][\uFE0F\uFE0E]?'), '✔'),
    # Crosses -> ✘ (\u2718)
    (re.compile(r'[\u274C\u274E\u2718][\uFE0F\uFE0E]?'), '✘'),
    # Stars -> ★ (\u2605)
    (re.compile(r'[\U0001F31F\U0001F4AB\u2605][\uFE0F\uFE0E]?'), '★'),
    # Pin / Fire -> 📍
    (re.compile(r'[\U0001F4CD\U0001F525][\uFE0F\uFE0E]?'), '📍'),
    # Rocket / Arrow -> ▶
    (re.compile(r'[\U0001F680\U0001F44D][\uFE0F\uFE0E]?'), '▶'),
]

def sanitize_emojis(text: str) -> str:
    """Convert high-surrogate emojis to clean BMP vector symbols supported by PDF engines."""
    if not text:
        return text
    for pattern, replacement in EMOJI_PATTERNS:
        text = pattern.sub(replacement, text)
    # Map any remaining high-surrogate characters to clean symbol
    text = re.sub(r'[\U00010000-\U0010FFFF]', '●', text)
    return text

def get_fa_unicode(tag) -> str:
    """Extract Unicode symbol for a FontAwesome <i> tag."""
    classes = tag.get('class', [])
    for cls in classes:
        if cls in FA_CLASS_TO_UNICODE:
            return FA_CLASS_TO_UNICODE[cls]
    return '●'


# ─────────────────────────────────────────────────────────────
# 1.  HTML → python-docx Document
# ─────────────────────────────────────────────────────────────

def _hex_to_rgb(hex_color: str) -> Optional[tuple]:
    """Convert #rrggbb or rgb(r,g,b) to (r,g,b) tuple."""
    if not hex_color:
        return None
    hex_color = hex_color.strip()
    m = re.match(r'rgb\((\d+),\s*(\d+),\s*(\d+)\)', hex_color)
    if m:
        return int(m.group(1)), int(m.group(2)), int(m.group(3))
    m = re.match(r'#?([0-9a-fA-F]{6})', hex_color)
    if m:
        h = m.group(1)
        return int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16)
    return None


def _pt_from_css(size_str: str, default: float = 11.0) -> float:
    """Parse CSS font-size string to points."""
    if not size_str:
        return default
    size_str = size_str.strip().lower()
    if size_str.endswith('px'):
        return float(size_str[:-2]) * 0.75
    if size_str.endswith('pt'):
        return float(size_str[:-2])
    if size_str.endswith('em'):
        return float(size_str[:-2]) * default
    return default


def _parse_inline_styles(tag) -> dict:
    """Extract inline CSS styles from a tag."""
    styles = {}
    style_str = tag.get('style', '') if hasattr(tag, 'get') else ''
    for part in style_str.split(';'):
        part = part.strip()
        if ':' in part:
            k, v = part.split(':', 1)
            styles[k.strip().lower()] = v.strip()
    return styles


def _apply_run_formatting(run, tag, inherited=None):
    """Apply bold/italic/underline/color/size to a docx Run."""
    if inherited is None:
        inherited = {}

    tag_name = tag.name if hasattr(tag, 'name') else ''
    styles = _parse_inline_styles(tag)
    merged = {**inherited, **styles}

    # Bold
    if tag_name in ('strong', 'b'):
        run.bold = True
    elif merged.get('font-weight', '') in ('bold', '700', '800', '900') or \
            (merged.get('font-weight', '0').isdigit() and int(merged.get('font-weight', '0')) >= 600):
        run.bold = True

    # Italic
    if tag_name in ('em', 'i'):
        run.italic = True
    elif merged.get('font-style', '') == 'italic':
        run.italic = True

    # Underline
    if tag_name == 'u' or 'underline' in merged.get('text-decoration', ''):
        run.underline = True

    # Color
    color_val = merged.get('color', '')
    rgb = _hex_to_rgb(color_val) if color_val else None
    if rgb:
        run.font.color.rgb = RGBColor(*rgb)

    # Font size
    font_size = merged.get('font-size', '')
    if font_size:
        pt = _pt_from_css(font_size)
        if pt > 0:
            run.font.size = Pt(pt)

    return run


def _add_inline_content(paragraph, element, inherited_styles=None):
    """Recursively add inline text content and icons to a paragraph."""
    if inherited_styles is None:
        inherited_styles = {}

    for child in element.children:
        if isinstance(child, NavigableString):
            text = str(child)
            if text:
                text = sanitize_emojis(text)
                run = paragraph.add_run(text)
                _apply_run_formatting(run, element, inherited_styles)
        elif isinstance(child, Tag):
            tag_name = child.name.lower() if child.name else ''
            if tag_name in ('script', 'style', 'svg'):
                continue
            if tag_name == 'br':
                paragraph.add_run('\n')
                continue
            # FontAwesome Icon support
            if tag_name in ('i', 'span') and any('fa' in c for c in child.get('class', [])):
                sym = get_fa_unicode(child)
                run = paragraph.add_run(sym + ' ')
                _apply_run_formatting(run, child, inherited_styles)
                continue
            if tag_name == 'img':
                continue
            if tag_name == 'a':
                link_text = sanitize_emojis(child.get_text())
                if link_text:
                    run = paragraph.add_run(link_text)
                    run.font.color.rgb = RGBColor(0x25, 0x63, 0xEB)
                    run.underline = True
                continue

            child_styles = _parse_inline_styles(child)
            merged = {**inherited_styles, **child_styles}
            _add_inline_content(paragraph, child, merged)


def _set_paragraph_format(para, tag_name: str, styles: dict):
    """Set paragraph-level formatting (alignment, spacing)."""
    alignment_map = {
        'center': WD_ALIGN_PARAGRAPH.CENTER,
        'right': WD_ALIGN_PARAGRAPH.RIGHT,
        'justify': WD_ALIGN_PARAGRAPH.JUSTIFY,
        'left': WD_ALIGN_PARAGRAPH.LEFT,
    }
    text_align = styles.get('text-align', '')
    if text_align in alignment_map:
        para.alignment = alignment_map[text_align]

    cls = styles.get('_class', '')
    if 'ql-align-center' in cls:
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif 'ql-align-right' in cls:
        para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    elif 'ql-align-justify' in cls:
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


def html_to_docx(html_content: str, doc_title: str = "SEO PDF Document") -> bytes:
    """
    Convert HTML string to DOCX bytes using python-docx.
    Preserves: Title (H1), headings, bold, italic, underline, color, font size,
               lists (ul/ol), icons/emojis, alignment, paragraphs, line spacing.
    """
    doc = Document()

    # --- Page setup: A4 ---
    section = doc.sections[0]
    section.page_width  = Cm(21.0)
    section.page_height = Cm(29.7)
    section.left_margin   = Cm(2.0)
    section.right_margin  = Cm(2.0)
    section.top_margin    = Cm(2.0)
    section.bottom_margin = Cm(2.0)

    # --- Default paragraph style ---
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)

    # --- Parse HTML ---
    soup = BeautifulSoup(html_content.replace('\uFEFF', ''), 'lxml')
    body = soup.body if soup.body else soup

    # Extract H1 for Document Title metadata
    h1_node = body.find('h1')
    effective_title = sanitize_emojis(h1_node.get_text().strip()) if (h1_node and h1_node.get_text().strip()) else doc_title

    # Document properties (for SEO indexing)
    core_props = doc.core_properties
    core_props.title = effective_title
    core_props.subject = effective_title
    core_props.keywords = effective_title
    core_props.author = 'SEO PDF Suite'

    heading_sizes = {'h1': 24, 'h2': 17, 'h3': 14, 'h4': 12, 'h5': 11, 'h6': 10}
    heading_colors = {
        'h1': RGBColor(0x0f, 0x17, 0x2a),
        'h2': RGBColor(0x1e, 0x29, 0x3b),
        'h3': RGBColor(0x33, 0x41, 0x55),
        'h4': RGBColor(0x47, 0x55, 0x69),
    }

    def process_node(node):
        if isinstance(node, NavigableString):
            return

        tag_name = (node.name or '').lower()
        inline_styles = _parse_inline_styles(node)
        inline_styles['_class'] = ' '.join(node.get('class', []))

        # ── Title (H1) & Headings (H2-H6) ──────────────────
        if tag_name in heading_sizes:
            para = doc.add_paragraph()
            para.paragraph_format.space_before = Pt(14 if tag_name == 'h1' else 10)
            para.paragraph_format.space_after  = Pt(8 if tag_name == 'h1' else 6)
            para.paragraph_format.line_spacing = Pt(heading_sizes[tag_name] * 1.3)
            _set_paragraph_format(para, tag_name, inline_styles)

            _add_inline_content(para, node)

            for r in para.runs:
                r.bold = True
                r.font.size = Pt(heading_sizes[tag_name])
                if tag_name in heading_colors and not r.font.color.rgb:
                    r.font.color.rgb = heading_colors[tag_name]
            return

        # ── Paragraph ─────────────────────────────────────
        if tag_name in ('p', 'div', 'section', 'article'):
            text_content = node.get_text(strip=False)
            if not text_content.strip():
                doc.add_paragraph()
                return
            para = doc.add_paragraph()
            para.paragraph_format.space_after  = Pt(7)
            para.paragraph_format.line_spacing = Pt(17)
            _set_paragraph_format(para, tag_name, inline_styles)
            _add_inline_content(para, node)
            return

        # ── Lists ──────────────────────────────────────────
        if tag_name in ('ul', 'ol'):
            for li in node.find_all('li', recursive=False):
                style_name = 'List Bullet' if tag_name == 'ul' else 'List Number'
                para = doc.add_paragraph(style=style_name)
                para.paragraph_format.space_after = Pt(4)
                para.paragraph_format.line_spacing = Pt(16)
                _add_inline_content(para, li)
            return

        # ── Horizontal Rule ────────────────────────────────
        if tag_name == 'hr':
            para = doc.add_paragraph()
            pPr = para._p.get_or_add_pPr()
            pBdr = OxmlElement('w:pBdr')
            bottom = OxmlElement('w:bottom')
            bottom.set(qn('w:val'), 'single')
            bottom.set(qn('w:sz'), '4')
            bottom.set(qn('w:space'), '1')
            bottom.set(qn('w:color'), 'CBD5E1')
            pBdr.append(bottom)
            pPr.append(pBdr)
            return

        # ── Blockquote ─────────────────────────────────────
        if tag_name == 'blockquote':
            para = doc.add_paragraph()
            para.paragraph_format.left_indent   = Cm(1.0)
            para.paragraph_format.space_after   = Pt(6)
            para.paragraph_format.line_spacing  = Pt(16)
            _add_inline_content(para, node)
            for r in para.runs:
                r.italic = True
                r.font.color.rgb = RGBColor(0x47, 0x55, 0x69)
            return

        # ── Tables ─────────────────────────────────────────
        if tag_name == 'table':
            rows = node.find_all('tr')
            if not rows:
                return
            cols = max(len(r.find_all(['td', 'th'])) for r in rows)
            if cols == 0:
                return
            tbl = doc.add_table(rows=len(rows), cols=cols)
            tbl.style = 'Table Grid'
            for r_idx, tr in enumerate(rows):
                cells = tr.find_all(['td', 'th'])
                for c_idx, cell in enumerate(cells):
                    if c_idx >= cols:
                        break
                    tbl_cell = tbl.rows[r_idx].cells[c_idx]
                    tbl_cell.text = sanitize_emojis(cell.get_text(strip=True))
                    if cell.name == 'th':
                        for r in tbl_cell.paragraphs[0].runs:
                            r.bold = True
            return

        for child in node.children:
            if isinstance(child, Tag):
                process_node(child)

    for node in body.children:
        if isinstance(node, Tag):
            process_node(node)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.read()


# ─────────────────────────────────────────────────────────────
# 2.  DOCX → PDF via LibreOffice Headless
# ─────────────────────────────────────────────────────────────

def _find_libreoffice() -> Optional[str]:
    """Find LibreOffice binary on macOS / Linux."""
    candidates = [
        '/Applications/LibreOffice.app/Contents/MacOS/soffice',
        '/usr/bin/libreoffice',
        '/usr/bin/soffice',
        '/usr/local/bin/libreoffice',
        shutil.which('libreoffice'),
        shutil.which('soffice'),
    ]
    for c in candidates:
        if c and Path(c).exists():
            return c
    return None


def docx_to_pdf_libreoffice(docx_bytes: bytes) -> bytes:
    soffice = _find_libreoffice()
    if not soffice:
        raise RuntimeError("LibreOffice not found.")

    with tempfile.TemporaryDirectory(prefix='seo_pdf_') as tmpdir:
        docx_path = Path(tmpdir) / 'document.docx'
        pdf_path  = Path(tmpdir) / 'document.pdf'

        docx_path.write_bytes(docx_bytes)

        cmd = [
            soffice,
            '--headless',
            '--norestore',
            '--nofirststartwizard',
            '--convert-to', 'pdf',
            '--outdir', tmpdir,
            str(docx_path)
        ]

        result = subprocess.run(
            cmd,
            capture_output=True,
            text=True,
            timeout=60
        )

        if result.returncode != 0:
            raise RuntimeError(f"LibreOffice conversion failed (code {result.returncode})")

        if not pdf_path.exists():
            raise RuntimeError(f"LibreOffice did not produce PDF.")

        return pdf_path.read_bytes()


# ─────────────────────────────────────────────────────────────
# 3.  PyMuPDF Optimization
# ─────────────────────────────────────────────────────────────

def optimize_pdf_pymupdf(pdf_bytes: bytes, doc_title: str = "SEO PDF Document") -> bytes:
    """
    Optimize PDF with PyMuPDF:
    - Set searchable metadata (title, subject, keywords → Google indexes these)
    - Compress streams
    Returns optimized PDF bytes.
    """
    doc = fitz.open(stream=pdf_bytes, filetype='pdf')

    doc.set_metadata({
        'title':    doc_title,
        'subject':  doc_title,
        'keywords': doc_title,
        'author':   'SEO PDF Suite',
        'creator':  'SEO PDF Suite via LibreOffice',
        'producer': 'PyMuPDF Optimizer',
    })

    out_buffer = io.BytesIO()
    doc.save(
        out_buffer,
        garbage=4,
        deflate=True,
        clean=True,
    )
    doc.close()

    out_buffer.seek(0)
    return out_buffer.read()


# ─────────────────────────────────────────────────────────────
# 4.  Layout Validation (automatic quality check)
# ─────────────────────────────────────────────────────────────

def validate_pdf_layout(pdf_bytes: bytes) -> dict:
    doc = fitz.open(stream=pdf_bytes, filetype='pdf')
    issues = []

    for page_num, page in enumerate(doc, 1):
        page_rect = page.rect
        blocks = page.get_text("blocks")

        for block in blocks:
            x0, y0, x1, y1, text, *_ = block
            if x1 > page_rect.width + 5:
                issues.append(f"Page {page_num}: Text overflow on right edge (x={x1:.0f}px)")
            if x0 < -5:
                issues.append(f"Page {page_num}: Text overflow on left edge (x={x0:.0f}px)")
            if y1 > page_rect.height + 5:
                issues.append(f"Page {page_num}: Text cut-off at bottom (y={y1:.0f}px)")

    doc.close()
    return {
        "pages": len(doc) if not doc.is_closed else 0,
        "valid": len(issues) == 0,
        "issues": issues
    }


# ─────────────────────────────────────────────────────────────
# 5.  ReportLab Fallback (with Arial Unicode & FontAwesome Support)
# ─────────────────────────────────────────────────────────────

REPORTLAB_FONT = 'Helvetica'
try:
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont
    
    font_candidates = [
        ('/Library/Fonts/Arial Unicode.ttf', 'ArialUnicode'),
        ('/System/Library/Fonts/Supplemental/Arial.ttf', 'ArialSys'),
        ('/Library/Fonts/Arial.ttf', 'ArialLib'),
    ]
    for font_path, font_name in font_candidates:
        if os.path.exists(font_path):
            pdfmetrics.registerFont(TTFont(font_name, font_path))
            REPORTLAB_FONT = font_name
            break
except Exception:
    pass


def _generate_pdf_reportlab(html_content: str, doc_title: str = "SEO PDF Document") -> bytes:
    """
    Fallback PDF generator using ReportLab.
    Produces 100% searchable, text-based PDF with FontAwesome icons & Emoji support.
    """
    from reportlab.lib.pagesizes import A4
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, HRFlowable
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib import colors

    buf = io.BytesIO()

    # Parse HTML & extract H1 for Title
    soup = BeautifulSoup(html_content.replace('\uFEFF', ''), 'lxml')
    body = soup.body if soup.body else soup
    h1_node = body.find('h1')
    effective_title = sanitize_emojis(h1_node.get_text().strip()) if (h1_node and h1_node.get_text().strip()) else doc_title

    doc = SimpleDocTemplate(
        buf, pagesize=A4,
        leftMargin=56, rightMargin=56, topMargin=56, bottomMargin=56,
        title=effective_title, author="SEO PDF Suite"
    )

    styles = getSampleStyleSheet()
    def S(name, parent, **kw):
        return ParagraphStyle(name, parent=styles[parent], **kw)

    title_style = S('DocTitle', 'Heading1', fontName=REPORTLAB_FONT, fontSize=24, leading=30, textColor=colors.HexColor('#0f172a'), spaceBefore=0,  spaceAfter=14)
    h2_style    = S('H2',       'Heading2', fontName=REPORTLAB_FONT, fontSize=17, leading=22, textColor=colors.HexColor('#1e293b'), spaceBefore=12, spaceAfter=8)
    h3_style    = S('H3',       'Heading3', fontName=REPORTLAB_FONT, fontSize=14, leading=18, textColor=colors.HexColor('#334155'), spaceBefore=8,  spaceAfter=6)
    h4_style    = S('H4',       'Normal',   fontName=REPORTLAB_FONT, fontSize=12, leading=16, textColor=colors.HexColor('#475569'), spaceBefore=6,  spaceAfter=4)
    bod_style   = S('Bod',      'Normal',   fontName=REPORTLAB_FONT, fontSize=11, leading=17, textColor=colors.HexColor('#1a1a1a'), spaceBefore=0,  spaceAfter=8)
    bul_style   = S('Bul',      'Normal',   fontName=REPORTLAB_FONT, fontSize=11, leading=16, textColor=colors.HexColor('#1a1a1a'), leftIndent=16, bulletIndent=4, spaceBefore=0, spaceAfter=4)

    story = []

    def clean_inline(el) -> str:
        """Convert element to ReportLab-safe markup string with icons."""
        parts = []
        for child in el.children:
            if isinstance(child, NavigableString):
                txt = str(child).replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
                txt = sanitize_emojis(txt)
                parts.append(txt)
            elif isinstance(child, Tag):
                t   = child.name.lower()
                # FontAwesome icon tag <i class="fa...">
                if t in ('i', 'span') and any('fa' in c for c in child.get('class', [])):
                    sym = get_fa_unicode(child)
                    parts.append(f'{sym} ')
                    continue

                sub = clean_inline(child)
                if t in ('strong', 'b'):
                    parts.append(f'<b>{sub}</b>')
                elif t in ('em', 'i'):
                    parts.append(f'<i>{sub}</i>')
                elif t == 'u':
                    parts.append(f'<u>{sub}</u>')
                elif t == 'a':
                    href = child.get('href', '')
                    parts.append(f'<link href="{href}"><u><font color="#2563eb">{sub}</font></u></link>')
                elif t == 'br':
                    parts.append('<br/>')
                elif t in ('span', 'font'):
                    style_str = child.get('style', '')
                    color_m   = re.search(r'color:\s*([^;]+)', style_str)
                    if color_m:
                        rgb = _hex_to_rgb(color_m.group(1).strip())
                        if rgb:
                            hex_c = '#{:02x}{:02x}{:02x}'.format(*rgb)
                            parts.append(f'<font color="{hex_c}">{sub}</font>')
                            continue
                    parts.append(sub)
                else:
                    parts.append(sub)
        return ''.join(parts)

    for node in body.children:
        if not isinstance(node, Tag):
            continue
        tag = node.name.lower()

        if tag == 'h1':
            story.append(Paragraph(f'<b>{clean_inline(node)}</b>', title_style))
        elif tag == 'h2':
            story.append(Paragraph(f'<b>{clean_inline(node)}</b>', h2_style))
        elif tag == 'h3':
            story.append(Paragraph(f'<b>{clean_inline(node)}</b>', h3_style))
        elif tag in ('h4', 'h5', 'h6'):
            story.append(Paragraph(f'<b>{clean_inline(node)}</b>', h4_style))
        elif tag in ('p', 'div', 'section', 'article'):
            txt = clean_inline(node).strip()
            if txt:
                story.append(Paragraph(txt, bod_style))
            else:
                story.append(Spacer(1, 8))
        elif tag == 'ul':
            for li in node.find_all('li', recursive=False):
                story.append(Paragraph(f'•  {clean_inline(li)}', bul_style))
        elif tag == 'ol':
            for idx2, li in enumerate(node.find_all('li', recursive=False), 1):
                story.append(Paragraph(f'{idx2}.  {clean_inline(li)}', bul_style))
        elif tag == 'blockquote':
            story.append(Paragraph(f'<i>{clean_inline(node)}</i>', bod_style))
        elif tag == 'hr':
            story.append(HRFlowable(width='100%', thickness=0.5, color=colors.HexColor('#CBD5E1')))
            story.append(Spacer(1, 6))

    if not story:
        txt = re.sub(r'<[^>]+>', ' ', html_content)
        txt = sanitize_emojis(txt)
        story.append(Paragraph(txt.strip(), bod_style))

    doc.build(story)
    pdf_bytes = buf.getvalue()
    buf.close()

    return optimize_pdf_pymupdf(pdf_bytes, effective_title)


# ─────────────────────────────────────────────────────────────
# 6.  Main Public API
# ─────────────────────────────────────────────────────────────

class LibreOfficePdfService:
    @classmethod
    def generate_pdf(cls, html_content: str, doc_title: str = "SEO PDF Document") -> bytes:
        if cls.libreoffice_available():
            docx_bytes = html_to_docx(html_content, doc_title)
            raw_pdf = docx_to_pdf_libreoffice(docx_bytes)
            return optimize_pdf_pymupdf(raw_pdf, doc_title)
        else:
            return _generate_pdf_reportlab(html_content, doc_title)

    @classmethod
    def libreoffice_available(cls) -> bool:
        return _find_libreoffice() is not None


class VectorPdfService:
    @classmethod
    def generate_vector_pdf_bytes(cls, html_content: str, doc_title: str = "SEO PDF Document") -> bytes:
        return LibreOfficePdfService.generate_pdf(html_content, doc_title)
